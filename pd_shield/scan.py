"""pd-scan: инвентаризация персональных данных в файлах.

CLI из шага 1, детекторы общие с masker (один источник истины).
Этим скриптом выполняется пункт 8 чек-листа перед сдачей:
«база и логи прогнаны скриптом инвентаризации ПД».

Использование:
    pd-scan ПУТЬ [ПУТЬ ...] -o отчёт.md [--config pd_config.json]

Форматы: txt, md, csv, json, docx, pdf, xlsx, sqlite/sqlite3/db.
Для docx/pdf/xlsx нужны зависимости: pip install "flameflow-pd-shield[scan]".
Режим инвентаризации перебирает, а не недобирает: одиночные
срабатывания NER включены, в отчёте они помечены отдельно.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sqlite3
from collections import defaultdict

from .detectors import (NameDictionary, merge_person_spans, ner_spans,
                        regex_spans, resolve)

SKIP_DIRS = {".git", "__pycache__", "venv", ".venv", "node_modules",
             ".pytest_cache"}

TYPE_TITLES = {
    "person": "ФИО",
    "phone": "телефон",
    "email": "email",
    "birthdate": "дата рождения",
    "address": "адрес",
}


# ---------------------------------------------------------------- извлечение

def text_from_txt(path):
    raw = open(path, "rb").read()
    for enc in ("utf-8", "cp1251"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def text_from_json(path):
    try:
        data = json.loads(text_from_txt(path))
    except Exception:
        return text_from_txt(path)
    parts = []

    def walk(x):
        if isinstance(x, dict):
            for k, v in x.items():
                parts.append(str(k))
                walk(v)
        elif isinstance(x, list):
            for v in x:
                walk(v)
        elif x is not None:
            parts.append(str(x))

    walk(data)
    return "\n".join(parts)


def text_from_docx(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in d.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                parts.append(p.text)
    return "\n".join(parts)


def text_from_pdf(path):
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
            page.close()  # без close pdfplumber копит кэш страниц до OOM
    return "\n".join(parts)


def text_from_xlsx(path):
    import openpyxl
    parts = []
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for ws in wb.worksheets:
        parts.append(ws.title)
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if cell is not None:
                    parts.append(str(cell))
    wb.close()
    return "\n".join(parts)


def text_from_sqlite(path):
    parts = []
    con = sqlite3.connect(f"file:{path}?mode=ro&immutable=1", uri=True)
    con.text_factory = lambda b: b.decode("utf-8", errors="ignore")
    cur = con.cursor()
    tables = [r[0] for r in cur.execute(
        "SELECT name FROM sqlite_master WHERE type='table'")]
    seen = set()
    for t in tables:
        try:
            cols = [r[1] for r in cur.execute(f'PRAGMA table_info("{t}")')]
            parts.append(f"[таблица {t}: {', '.join(cols)}]")
            for row in cur.execute(f'SELECT * FROM "{t}"'):
                for cell in row:
                    if isinstance(cell, str) and len(cell) > 1 and cell not in seen:
                        seen.add(cell)
                        parts.append(cell)
        except sqlite3.Error as e:
            parts.append(f"[таблица {t}: ошибка чтения {e}]")
    con.close()
    return "\n".join(parts)


EXTRACTORS = {
    ".txt": text_from_txt, ".md": text_from_txt, ".csv": text_from_txt,
    ".json": text_from_json, ".docx": text_from_docx, ".pdf": text_from_pdf,
    ".xlsx": text_from_xlsx, ".sqlite": text_from_sqlite,
    ".sqlite3": text_from_sqlite, ".db": text_from_sqlite,
}


def iter_files(paths):
    for p in paths:
        if os.path.isfile(p):
            yield p
        else:
            for root, dirs, files in os.walk(p):
                dirs[:] = [d for d in dirs if d not in SKIP_DIRS]
                for f in sorted(files):
                    yield os.path.join(root, f)


# ---------------------------------------------------------------- скан

def scan_text(text: str, dictionary: NameDictionary | None = None):
    """Находки инвентаризации: (тип-для-отчёта, значение)."""
    spans = regex_spans(text)
    if dictionary is not None:
        spans += dictionary.spans(text)
    spans += ner_spans(text, known_surnames=(dictionary.exact
                                             if dictionary else set()),
                       include_single=True)
    spans = merge_person_spans(
        text, spans, protected=dictionary.exact if dictionary else None)
    spans = resolve(spans, list(TYPE_TITLES))
    out = []
    for s in spans:
        title = TYPE_TITLES[s.type]
        if s.type == "person" and " " not in s.text.strip():
            title = "ФИО (одно слово, проверить)"
        out.append((title, re.sub(r"\s+", " ", s.text).strip(" ,;:.")))
    return out


def scan(paths, base=None, dictionary=None):
    findings = defaultdict(list)
    errors, scanned, skipped = [], [], []
    for path in iter_files(paths):
        ext = os.path.splitext(path)[1].lower()
        rel = os.path.relpath(path, base) if base else path
        name_hits = [(t + " в имени файла", v)
                     for t, v in scan_text(os.path.basename(path), dictionary)]
        extractor = EXTRACTORS.get(ext)
        if extractor is None:
            skipped.append(rel)
            if name_hits:
                findings[rel].extend(name_hits)
            continue
        scanned.append(rel)
        findings[rel].extend(name_hits)
        try:
            text = extractor(path)
        except Exception as e:
            errors.append((rel, str(e)))
            continue
        seen = set()
        for t, v in scan_text(text, dictionary):
            key = (t, v.lower())
            if key not in seen:
                seen.add(key)
                findings[rel].append((t, v))
    return findings, scanned, skipped, errors


def write_report(out_path, findings, scanned, skipped, errors, title_paths):
    lines = ["# Инвентаризация персональных данных", ""]
    total = sum(len(v) for v in findings.values())
    files_with = sorted(f for f, v in findings.items() if v)
    lines.append(f"Проверено: {', '.join(title_paths)}")
    lines.append(f"Файлов прочитано: {len(scanned)}, пропущено по формату: "
                 f"{len(skipped)}, ошибок чтения: {len(errors)}")
    lines.append(f"Файлов с находками: {len(files_with)}, всего находок: {total}")
    lines.append("")
    for f in files_with:
        lines += [f"## {f}", "", "| Тип | Находка |", "|---|---|"]
        for t, v in sorted(findings[f]):
            esc = v.replace("|", "\\|")
            lines.append(f"| {t} | {esc} |")
        lines.append("")
    if errors:
        lines += ["## Ошибки чтения", ""]
        lines += [f"- {f}: {e}" for f, e in errors] + [""]
    if skipped:
        lines += ["## Пропущено по формату", ""]
        lines += [f"- {f}" for f in sorted(skipped)] + [""]
    with open(out_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))


def main():
    ap = argparse.ArgumentParser(description="Инвентаризация ПД")
    ap.add_argument("paths", nargs="+")
    ap.add_argument("-o", "--out", default="pd_report.md")
    ap.add_argument("--base", default=None,
                    help="база для относительных путей в отчёте")
    ap.add_argument("--config", default=None,
                    help="pd_config.json проекта: словарь имён участвует в скане")
    args = ap.parse_args()

    dictionary = None
    if args.config:
        from .config import ShieldConfig
        dictionary = NameDictionary(ShieldConfig.from_json(args.config).names)

    findings, scanned, skipped, errors = scan(args.paths, base=args.base,
                                              dictionary=dictionary)
    write_report(args.out, findings, scanned, skipped, errors, args.paths)
    total = sum(len(v) for v in findings.values())
    print(f"файлов прочитано: {len(scanned)}, находок: {total}, "
          f"ошибок: {len(errors)}, отчёт: {args.out}")


if __name__ == "__main__":
    main()
